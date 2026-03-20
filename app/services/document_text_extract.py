"""
Извлечение текста из загруженных документов (analyze first pass).
PDF — только текстовый слой (без OCR). DOCX — параграфы и таблицы.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import List, Optional

from docx import Document
from pypdf import PdfReader

# Лимит размера файла (байты). Согласуйте с лимитами Telegram / хостинга.
MAX_UPLOAD_BYTES = 15 * 1024 * 1024

# Макс. длина текста тела документа после извлечения / вставки (символы).
# Согласовано порядка с OPENAI_MAX_INPUT_CHARS + запас под промпт/RAG.
MAX_EXTRACTED_TEXT_CHARS = 150_000

WARN_PDF_NO_TEXT_LAYER = "pdf_no_text_layer"
WARN_DOCX_EMPTY = "docx_empty"
WARN_TEXT_TRUNCATED = "text_truncated_length"


@dataclass
class ExtractResult:
    text: str
    ok: bool
    warnings: List[str] = field(default_factory=list)
    error_message: Optional[str] = None


def _normalize_text(raw: str) -> str:
    if not raw:
        return ""
    lines = [ln.strip() for ln in raw.splitlines()]
    parts: List[str] = []
    buf: List[str] = []
    for ln in lines:
        if ln:
            buf.append(ln)
        else:
            if buf:
                parts.append(" ".join(buf))
                buf = []
    if buf:
        parts.append(" ".join(buf))
    return "\n\n".join(parts).strip()


def _docx_extract(data: bytes) -> tuple[str, List[str]]:
    warnings: List[str] = []
    doc = Document(BytesIO(data))
    chunks: List[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            chunks.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    text = _normalize_text("\n\n".join(chunks))
    if not text:
        warnings.append(WARN_DOCX_EMPTY)
    return text, warnings


def _pdf_extract(data: bytes) -> tuple[str, List[str]]:
    warnings: List[str] = []
    reader = PdfReader(BytesIO(data))
    parts: List[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
        except Exception:
            t = None
        if t and t.strip():
            parts.append(t.strip())
    text = _normalize_text("\n\n".join(parts))
    if not text:
        warnings.append(WARN_PDF_NO_TEXT_LAYER)
    return text, warnings


def clamp_document_text_for_analyze(
    text: str,
    *,
    warnings: Optional[List[str]] = None,
) -> tuple[str, List[str]]:
    """
    Единая обрезка длинного тела документа для сценария analyze (файл и вставка).
    """
    w = list(warnings or [])
    if len(text) <= MAX_EXTRACTED_TEXT_CHARS:
        return text, w
    w.append(WARN_TEXT_TRUNCATED)
    return text[:MAX_EXTRACTED_TEXT_CHARS], w


def _resolve_kind(filename: Optional[str], mime_type: Optional[str]) -> Optional[str]:
    if filename:
        ext = Path(filename).suffix.lower()
        if ext == ".doc":
            return "doc_legacy"
        if ext == ".pdf":
            return "pdf"
        if ext == ".docx":
            return "docx"
    mt = (mime_type or "").lower().split(";")[0].strip()
    if mt == "application/pdf":
        return "pdf"
    if mt == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx"
    if mt == "application/msword":
        return "doc_legacy"
    return None


def extract_uploaded_document(
    *,
    filename: Optional[str],
    mime_type: Optional[str],
    data: bytes,
) -> ExtractResult:
    if len(data) > MAX_UPLOAD_BYTES:
        return ExtractResult(
            text="",
            ok=False,
            error_message=(
                f"Файл слишком большой (макс. {MAX_UPLOAD_BYTES // (1024 * 1024)} МБ). "
                "Сократи объём или вставь текст вручную."
            ),
        )

    kind = _resolve_kind(filename, mime_type)
    if kind == "doc_legacy":
        return ExtractResult(
            text="",
            ok=False,
            error_message=(
                "Формат <b>.doc</b> (старый Word) не поддерживается.\n"
                "Сохрани документ как <b>DOCX</b> или пришли <b>PDF</b> с текстовым слоем, либо вставь текст."
            ),
        )

    if kind not in ("pdf", "docx"):
        return ExtractResult(
            text="",
            ok=False,
            error_message=(
                "Формат файла не поддерживается.\n"
                "Можно: <b>PDF</b> (с текстовым слоем) или <b>DOCX</b>."
            ),
        )

    try:
        if kind == "docx":
            text, warnings = _docx_extract(data)
        else:
            text, warnings = _pdf_extract(data)
    except Exception as e:  # noqa: BLE001 — показываем пользователю общее сообщение
        return ExtractResult(
            text="",
            ok=False,
            error_message=(
                "Не удалось прочитать файл.\n"
                f"Причина: <code>{e}</code>\n"
                "Попробуй другой файл или вставь текст."
            ),
        )

    if not text.strip():
        msg = (
            "<b>Текст не найден.</b>\n"
            "Для PDF: возможно, скан без текстового слоя (OCR не поддерживается).\n"
            "Пришли DOCX, PDF с текстом или вставь текст вручную."
        )
        if kind == "docx" and WARN_DOCX_EMPTY in warnings:
            msg = (
                "<b>Текст не найден.</b>\n"
                "В DOCX нет извлекаемого текста (пустой файл или нестандартная структура).\n"
                "Проверь файл или вставь текст."
            )
        return ExtractResult(text="", ok=False, warnings=warnings, error_message=msg)

    text, warnings = clamp_document_text_for_analyze(text, warnings=warnings)
    return ExtractResult(text=text, ok=True, warnings=warnings)

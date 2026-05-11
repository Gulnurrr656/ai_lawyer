from pathlib import Path
from typing import Optional

from docx import Document


# =====================================================
# EXPORT UTILITIES — STABLE VERSION (DOCX ONLY)
# =====================================================
# ✔ Реальный DOCX (python-docx)
# ✔ Без PDF
# =====================================================


EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def save_docx(text: str, filename: str) -> str:
    """
    Сохраняет корректный DOCX-файл.
    """
    path = EXPORT_DIR / f"{filename}.docx"

    doc = Document()

    for block in text.split("\n\n"):
        p = doc.add_paragraph()
        p.add_run(block)

    doc.save(path)
    return str(path.resolve())
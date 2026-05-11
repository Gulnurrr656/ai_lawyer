"""DOCX (Word) export for ``GeneratedDocument``.

Renders a clean Word document with:

- Title (Heading 1)
- Each section as Heading 2 + body paragraphs
- Legal basis map table-like list
- Conflicts / defense / procedure blocks
- Appendices, used docs, missing facts, warnings

Cyrillic is preserved (UTF-8 throughout — python-docx handles it natively).

Public API:
- :func:`render_document_docx(document, output_path) -> str`
- :func:`is_docx_supported() -> bool`
"""

from __future__ import annotations

import os
from typing import Iterable, List

from app.shared.legal_engine_v2.document_schemas import GeneratedDocument


def is_docx_supported() -> bool:
    """Return True if python-docx is available in the runtime."""
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def _split_paragraphs(text: str) -> List[str]:
    if not text:
        return []
    out: List[str] = []
    for raw in text.splitlines():
        s = raw.rstrip()
        out.append(s)
    return out


def render_document_docx(document: GeneratedDocument, output_path: str) -> str:
    """Render ``document`` as a DOCX file at ``output_path``.

    Returns the absolute path of the saved file.
    Raises :class:`RuntimeError` if python-docx is not installed.
    """
    if not is_docx_supported():
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    # ---- Default font / size --------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # ---- Title ----------------------------------------------------------
    title = (document.title or document.doc_type or "ДОКУМЕНТ").strip()
    h_title = doc.add_heading(title, level=0)
    h_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Sections -------------------------------------------------------
    for i, s in enumerate(document.sections, 1):
        section_title = (s.title or s.key or "Раздел").strip()
        doc.add_heading(f"{i}. {section_title}", level=1)
        body = (s.draft_text or "").strip()
        if not body:
            doc.add_paragraph(
                "[УКАЗАТЬ СОДЕРЖАНИЕ РАЗДЕЛА]"
            )
            continue
        for para in _split_paragraphs(body):
            doc.add_paragraph(para)

    # ---- Legal basis map -----------------------------------------------
    if document.legal_basis_map:
        from app.shared.legal_engine_v2.legal_hierarchy import LEGAL_BASIS_MAP_KEYS
        from app.shared.legal_engine_v2.document_writer import _LEGAL_BASIS_LABELS

        added_header = False
        for slot in LEGAL_BASIS_MAP_KEYS:
            docs = document.legal_basis_map.get(slot) or []
            if not docs:
                continue
            if not added_header:
                doc.add_heading("ПРАВОВАЯ КАРТА (LEGAL BASIS MAP)", level=1)
                doc.add_paragraph(
                    "Иерархия источников права РК: Конституция → Кодексы → "
                    "Законы → НП ВС/КС → подзаконные акты."
                )
                added_header = True
            doc.add_heading(_LEGAL_BASIS_LABELS.get(slot, slot), level=2)
            for d in docs:
                line = d.article_label()
                src = (getattr(d, "source", "") or "").strip()
                if src and src not in line:
                    line = f"{line} ({src})"
                doc.add_paragraph(f"— {line}")

    # ---- Conflicts ------------------------------------------------------
    if document.conflicts_analysis:
        doc.add_heading("Иерархия применимых норм и разрешение возможных противоречий", level=1)
        for c in document.conflicts_analysis:
            line = f"— Тип: {c.get('conflict_type', '')}"
            if c.get("issue"):
                line += f"; вопрос: {c['issue']}"
            doc.add_paragraph(line)
            if c.get("resolution_rule"):
                doc.add_paragraph(f"  Правило разрешения: {c['resolution_rule']}")
            if c.get("recommended_argument"):
                doc.add_paragraph(f"  Рекомендуемый аргумент: {c['recommended_argument']}")
            for w in (c.get("warnings") or []):
                doc.add_paragraph(f"  Предупреждение: {w}")
            if not c.get("resolved", True):
                doc.add_paragraph("  Статус: НЕРАЗРЕШЁН — требуется ручная проверка.")

    # ---- Defense -------------------------------------------------------
    if document.defense_arguments:
        doc.add_heading("Позиция защиты", level=1)
        for a in document.defense_arguments:
            doc.add_paragraph(f"— Вопрос: {a.get('issue', '')}")
            if a.get("taxpayer_or_client_position"):
                doc.add_paragraph(f"  Позиция клиента: {a['taxpayer_or_client_position']}")
            if a.get("authority_position"):
                doc.add_paragraph(f"  Позиция оппонента: {a['authority_position']}")
            if a.get("special_norm"):
                doc.add_paragraph(f"  Специальная норма: {a['special_norm']}")
            if a.get("procedural_guarantees"):
                doc.add_paragraph(
                    "  Процессуальные гарантии: "
                    + "; ".join(a["procedural_guarantees"])
                )
            if a.get("attack_points"):
                doc.add_paragraph(
                    "  Точки атаки: " + "; ".join(a["attack_points"])
                )
            if a.get("evidence_needed"):
                doc.add_paragraph(
                    "  Какие доказательства собрать: "
                    + "; ".join(a["evidence_needed"])
                )
            if a.get("recommended_action"):
                doc.add_paragraph(f"  Рекомендация: {a['recommended_action']}")

    # ---- Procedure guide -----------------------------------------------
    if document.procedure_guide:
        doc.add_heading("Порядок подачи и практические действия", level=1)
        guide = document.procedure_guide
        route = (guide.get("filing_route") or {}) if isinstance(guide, dict) else {}
        doc.add_paragraph(f"Маршрут: {route.get('route_type', '')}")
        if route.get("authority_name"):
            doc.add_paragraph(f"Орган / адресат: {route['authority_name']}")
        if route.get("platform_name"):
            doc.add_paragraph(f"Платформа: {route['platform_name']}")
        doc.add_paragraph(f"Требуется ЭЦП: {'да' if route.get('requires_eds') else 'нет'}")
        if route.get("file_format"):
            doc.add_paragraph("Формат файла: " + ", ".join(route["file_format"]))
        if route.get("submission_method"):
            doc.add_paragraph("Способы подачи:")
            for m in route["submission_method"]:
                doc.add_paragraph(f"  • {m}")
        if route.get("required_documents"):
            doc.add_paragraph("Документы / приложения:")
            for d_ in route["required_documents"]:
                doc.add_paragraph(f"  • {d_}")
        if route.get("state_fee_or_costs"):
            doc.add_paragraph("Госпошлина / расходы:")
            for f_ in route["state_fee_or_costs"]:
                doc.add_paragraph(f"  • {f_}")
        if route.get("next_steps_after_submission"):
            doc.add_paragraph("После подачи:")
            for n_ in route["next_steps_after_submission"]:
                doc.add_paragraph(f"  • {n_}")
        if route.get("risks"):
            doc.add_paragraph("Риски:")
            for r_ in route["risks"]:
                doc.add_paragraph(f"  • {r_}")
        if guide.get("steps"):
            doc.add_heading("Шаги", level=2)
            for s_ in guide["steps"]:
                doc.add_paragraph(f"{s_.get('step_number', '?')}. {s_.get('title', '')}")
                if s_.get("description"):
                    doc.add_paragraph(f"   {s_['description']}")
        if guide.get("missing_procedure_facts"):
            doc.add_heading("Уточнить перед подачей", level=2)
            for m_ in guide["missing_procedure_facts"]:
                doc.add_paragraph(f"  • {m_}")

    # ---- Used docs -----------------------------------------------------
    if document.used_docs:
        doc.add_heading("Использованные источники (RAG)", level=1)
        for d in document.used_docs[:50]:
            line = d.article_label()
            src = (getattr(d, "source", "") or "").strip()
            if src and src not in line:
                line = f"{line} ({src})"
            bucket = (getattr(d, "bucket", "") or "").strip()
            if bucket:
                line = f"[{bucket}] {line}"
            doc.add_paragraph(f"— {line}")

    # ---- Missing facts -------------------------------------------------
    if document.missing_facts:
        doc.add_heading("Пробелы данных (MISSING FACTS)", level=1)
        for m in document.missing_facts:
            doc.add_paragraph(f"— {m}")

    # ---- Warnings ------------------------------------------------------
    if document.warnings:
        doc.add_heading("Предупреждения", level=1)
        for w in document.warnings:
            doc.add_paragraph(f"— {w}")

    # ---- Save ----------------------------------------------------------
    abs_path = os.path.abspath(output_path)
    parent = os.path.dirname(abs_path)
    if parent and not os.path.isdir(parent):
        os.makedirs(parent, exist_ok=True)
    doc.save(abs_path)
    return abs_path


__all__ = [
    "is_docx_supported",
    "render_document_docx",
]
